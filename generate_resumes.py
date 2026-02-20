from __future__ import annotations

import json
import re
import shutil
import subprocess
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "resumes"
CONFIG_PATH = ROOT / "resume_config.json"

Variant = str
SectionBody = Union[str, List[str]]
Section = Tuple[str, SectionBody]

ALLOWED_DEFAULT_FORMATS = ["classic", "madakor"]
ALLOWED_DEFAULT_VARIANTS = ["general", "python-backend", "cybersecurity", "ai-architect", "grc-analyst"]


# =========================
# SAFE SLUGS + DYNAMIC LISTS
# =========================
def safe_slug(name: str, default: str = "classic") -> str:
    """
    Convert arbitrary strings into safe folder/keys:
    - lowercase
    - spaces -> hyphen
    - remove unsafe characters
    - collapse multiple hyphens
    """
    if not name:
        return default
    s = str(name).strip().lower()
    s = s.replace(" ", "-")
    s = re.sub(r"[^a-z0-9._-]+", "", s)
    s = re.sub(r"-{2,}", "-", s).strip("-")
    return s or default


def get_formats(cfg: "ResumeConfig") -> List[str]:
    formats = getattr(cfg, "formats", None)
    if isinstance(formats, list) and formats:
        cleaned = [safe_slug(x, default="classic") for x in formats if str(x).strip()]
        return sorted(list(dict.fromkeys(cleaned)))
    return [safe_slug(x, default="classic") for x in ALLOWED_DEFAULT_FORMATS]


def get_variants(cfg: "ResumeConfig") -> List[str]:
    variants = getattr(cfg, "variants", None)
    if isinstance(variants, list) and variants:
        cleaned = [safe_slug(x, default="general") for x in variants if str(x).strip()]
        return sorted(list(dict.fromkeys(cleaned)))
    return [safe_slug(x, default="general") for x in ALLOWED_DEFAULT_VARIANTS]


def format_folder_name(cfg: "ResumeConfig") -> str:
    # Always safe and consistent
    return safe_slug(getattr(cfg, "resume_format", "classic"), default="classic")


# =========================
# CONFIG
# =========================
@dataclass
class ResumeConfig:
    name: str = "BRIAN HANNIGAN"
    phone: str = "(862) 243-1177"
    email: str = "jobshannigan@gmail.com"
    linkedin: str = "linkedin.com/in/brianjhannigan"
    github: str = "github.com/brianhannigan"
    location: str = "Remote | 07876"
    libreoffice_path: str = ""
    resume_format: str = "classic"

    # NEW: user-extensible lists (saved in resume_config.json)
    formats: List[str] = None
    variants: List[str] = None


def load_config() -> ResumeConfig:
    if CONFIG_PATH.exists():
        data = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
        base = asdict(ResumeConfig())
        base.update(data)
        return ResumeConfig(**base)
    cfg = ResumeConfig()
    save_config(cfg)
    return cfg


def save_config(cfg: ResumeConfig) -> None:
    CONFIG_PATH.write_text(json.dumps(asdict(cfg), indent=2), encoding="utf-8")


# =========================
# DOC HELPERS
# =========================
def set_doc_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_center_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_section(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11.5)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)


def add_bullets(doc: Document, bullets: List[str]) -> None:
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(10.5)


def add_bold_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10.75)


def save_txt(path: Path, lines: List[str]) -> None:
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


# =========================
# LIBREOFFICE EXPORT
# =========================
def find_soffice(cfg: ResumeConfig) -> Optional[str]:
    if cfg.libreoffice_path and Path(cfg.libreoffice_path).exists():
        return cfg.libreoffice_path

    soffice = shutil.which("soffice")
    if soffice:
        return soffice

    default = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")
    if default.exists():
        return str(default)

    return None


def convert_with_libreoffice(docx_path: Path, fmt: str, cfg: ResumeConfig) -> Path:
    soffice = find_soffice(cfg)
    if not soffice:
        raise FileNotFoundError(
            "LibreOffice (soffice.exe) not found. Install LibreOffice or set libreoffice_path in resume_config.json"
        )

    cmd = [
        soffice,
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        "--convert-to",
        fmt,
        "--outdir",
        str(docx_path.parent),
        str(docx_path),
    ]
    subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

    out = docx_path.parent / f"{docx_path.stem}.{fmt}"
    if not out.exists():
        raise RuntimeError(f"LibreOffice conversion ran but {fmt.upper()} was not created.")
    return out


# =========================
# OUTPUT STRUCTURE
# =========================
def ensure_folders(cfg: ResumeConfig) -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    variants = get_variants(cfg)
    formats = get_formats(cfg)
    for v in variants:
        for f in formats:
            (OUT / v / f).mkdir(parents=True, exist_ok=True)


def variant_slug(variant: Variant) -> str:
    return f"brian-hannigan-{variant}"


# =========================
# SHARED CONTENT BLOCKS
# =========================
def education_lines_short() -> List[str]:
    return [
        "Master’s — New Jersey City University",
        "B.S. — University of Delaware",
        "Programming Diploma — Chubb Institute",
    ]


def certification_lines() -> List[str]:
    return [
        "CompTIA Security+",
        "InfraGard Member",
        "SOC Core Skills",
        "Active Defense & Cyber Deception",
        "Red Team Summit",
        "Hacking Humans: Social Engineering & Elicitation",
    ]


def logn_pacific_block() -> Dict:
    bullets = [
        "Conducted vulnerability scans and provided detailed remediation reports, contributing to a 100% reduction in critical, 90% reduction in high, and 76% reduction in medium vulnerabilities.",
        "Performed vulnerability assessments and risk prioritization using Tenable across Windows and Linux environments.",
        "Executed secure configurations and compliance audits (DISA STIG) using Tenable to meet industry standards.",
        "Automated remediation processes and STIG implementations using PowerShell to address critical vulnerabilities.",
        "Performed threat hunting using EDR, identifying indicators of brute force attacks, data exfiltration, and ransomware activity.",
        "Designed and tested advanced threat hunting scenarios for incident response tabletop exercises.",
        "Developed custom detection rules in Microsoft Defender for Endpoint to automate isolation and investigation of compromised systems.",
        "Reduced brute force incidents by 100% by implementing inbound NSG/firewall rules to limit Internet exposure.",
        "Created Microsoft Sentinel dashboards to monitor logon failures and malicious traffic using threat intelligence.",
        "Utilized KQL to query logs within SIEM and EDR platforms for investigation and detection purposes.",
    ]
    return {
        "company": "Log(N) Pacific",
        "dates": "2025 – Present",
        "title": "Cyber Security Support Analyst (Vulnerability Management & SecOps Intern)",
        "areas": [
            ("Vulnerability Management", bullets[:4]),
            ("Security Operations", bullets[4:]),
        ],
    }


def _known_variants() -> List[str]:
    return ["general", "python-backend", "cybersecurity", "ai-architect", "grc-analyst"]


def _fallback_variant(variant: Variant) -> Variant:
    # Safety: if UI adds a new variant without content, we don’t crash
    return variant if variant in _known_variants() else "general"


def utrs_block_for_variant(variant: Variant) -> Dict:
    variant = _fallback_variant(variant)

    base = {
        "company": "UTRS, Inc",
        "dates": "2011 – 2025",
        "title": "Technical Systems Engineer / Trainer / Systems Architect",
        "areas": [],
    }

    systems_delivery = [
        "Designed and delivered simulation training ecosystems for U.S. Government/DoD programs across multiple sites.",
        "Built and sustained operator tooling, automation workflows, and repeatable deployment procedures.",
        "Produced SOPs, technical documentation, and handoff-ready training materials for long-term sustainment.",
        "Partnered with stakeholders to translate requirements into deliverable system outcomes.",
        "Improved operational reliability through disciplined troubleshooting and root cause analysis.",
    ]

    secure_deploy = [
        "Applied secure configuration practices and compliance expectations (DISA STIG alignment awareness) in deployments.",
        "Supported remediation and reliability improvements through documentation-driven sustainment workflows.",
    ]

    if variant == "python-backend":
        base["areas"] = [
            ("Backend / Automation", [
                "Developed automation utilities and service-style tooling to support deployed systems and workflows.",
                "Implemented structured logging and debugging patterns to improve maintainability and reliability.",
                "Used Git workflows and documentation-first delivery for sustainment and repeatability.",
            ]),
            ("Systems Delivery", systems_delivery),
        ]
    elif variant == "cybersecurity":
        base["areas"] = [
            ("Secure Deployment & Compliance", secure_deploy + [
                "Supported vulnerability and remediation workflows in coordination with operational needs and constraints.",
            ]),
            ("Operational Troubleshooting", [
                "Investigated and resolved system issues across Windows-based environments to reduce downtime and risk.",
                "Maintained clear written documentation to support auditability and long-term sustainment.",
            ]),
        ]
    elif variant == "ai-architect":
        base["areas"] = [
            ("Systems Architecture", systems_delivery),
            ("Reliability & Sustainment", [
                "Designed repeatable deployment patterns and documentation practices to support adoption and sustainment.",
                "Emphasized deterministic behavior, logging, and supportability across delivered systems.",
            ]),
        ]
    elif variant == "grc-analyst":
        base["areas"] = [
            ("Documentation & Evidence Mindset", [
                "Produced SOPs, deployment artifacts, and sustainment materials supporting audit readiness.",
                "Organized system documentation and operational records to support compliance-aligned reporting.",
                "Tracked issues and follow-ups using structured templates and clear written status updates.",
            ]),
            ("Secure Configuration Support", secure_deploy),
        ]
    else:  # general
        base["areas"] = [
            ("Systems Engineering", systems_delivery),
            ("Secure Deployment Support", secure_deploy),
        ]

    return base


def consulting_block_for_variant(variant: Variant) -> Dict:
    variant = _fallback_variant(variant)

    base = {
        "company": "Independent Consulting",
        "dates": "Various",
        "title": "Software Engineer / Security & Automation Consultant",
        "areas": [],
    }

    common = [
        "Delivered client-facing software and automation solutions with strong documentation and handoff discipline.",
        "Built repeatable workflows, templates, and operational guides to reduce friction and improve consistency.",
        "Partnered directly with stakeholders; translated ambiguous needs into clear technical deliverables.",
    ]

    if variant == "python-backend":
        base["areas"] = [
            ("Python / Backend Delivery", [
                "Built Python utilities and API-driven workflows to automate business processes and data handling.",
                "Implemented structured logging, validation, and error handling patterns for supportability.",
            ] + common),
        ]
    elif variant == "cybersecurity":
        base["areas"] = [
            ("Security & Compliance Support", [
                "Supported compliance-aligned documentation and secure configuration expectations in delivered solutions.",
                "Created trackers and artifacts to support evidence organization and remediation follow-ups.",
            ] + common),
        ]
    elif variant == "grc-analyst":
        base["areas"] = [
            ("GRC / Audit Support", [
                "Created evidence-friendly documentation, trackers, and structured reporting templates for assessments.",
                "Supported remediation planning and follow-up workflows to drive closure and risk reduction.",
            ] + common),
        ]
    elif variant == "ai-architect":
        base["areas"] = [
            ("Automation & AI-Adjacent Workflows", [
                "Designed workflow automation patterns suitable for offline-first and constrained environments.",
                "Emphasized reliability, observability, and maintainable modular boundaries.",
            ] + common),
        ]
    else:
        base["areas"] = [
            ("Delivery & Enablement", common),
        ]

    return base


def translator_project_for_variant(variant: Variant) -> Dict:
    variant = _fallback_variant(variant)

    base_bullets = [
        "Designed and implemented Python-based REST services supporting OCR and AI translation workflows.",
        "Built HTTP endpoints with structured JSON request/response handling and robust validation.",
        "Integrated Tesseract OCR and local translation engines into modular, testable service boundaries.",
        "Implemented logging, error handling, and process monitoring to improve reliability and debuggability.",
        "Used Git for version control and maintained clean repo organization and documentation.",
    ]

    if variant == "cybersecurity":
        return {
            "title": "Threat Hunting + KQL Projects (Plus Portable Translator)",
            "source": "github.com/brianhannigan/kql-threathunting-beginner-guide",
            "platforms": "KQL, SIEM Concepts (Sentinel), EDR Concepts (MDE), GitHub, Windows",
            "bullets": [
                "Developed KQL-style hunting queries and investigation patterns for common attacker behaviors.",
                "Organized portfolio assets to demonstrate detection engineering approach and documentation discipline.",
            ] + base_bullets[:2],
        }

    return {
        "title": "Portable Translator — Offline OCR + Translation Services",
        "source": "github.com/brianhannigan/translator",
        "platforms": "Python, REST APIs, Tesseract OCR, Local Translation Engines, Git, Windows/Linux",
        "bullets": base_bullets,
    }


# =========================
# CLASSIC FORMAT CONTENT
# =========================
def classic_content_variant(variant: Variant) -> Tuple[str, List[str], List[Section]]:
    variant = _fallback_variant(variant)
    translator_lines = translator_project_for_variant("python-backend")["bullets"]

    if variant == "general":
        summary = (
            "Senior Software Engineer, Technical Systems Architect, and security-focused builder with 15+ years delivering "
            "secure, scalable systems across mission-critical environments. Experienced in backend services, automation, "
            "API-driven workflows, and deployment standardization."
        )
        skills = [
            "Languages: Python, C#, SQL, Bash",
            "Backend: REST APIs, JSON, HTTP, Service-Oriented Design",
            "Databases: SQLite, SQL fundamentals, relational modeling",
            "Tools: Git, Docker, Windows, Linux",
            "Security: Hardening, compliance-aligned engineering, vulnerability scanning concepts",
            "Strengths: Debugging, documentation, requirements translation, cross-functional delivery",
        ]
        sections: List[Section] = [
            ("Experience", [
                "Log(N) Pacific — Cyber Security Support Analyst (Vulnerability Management & SecOps Intern) (2025–Present)",
                "Vulnerability Management: Tenable scans, risk prioritization, DISA STIG audits, PowerShell automation.",
                "Security Operations: EDR threat hunting, MDE detections, Sentinel dashboards, KQL investigations, NSG/firewall hardening.",
                "",
                "UTRS, Inc — Technical Systems Engineer / Trainer / Systems Architect (2011–2025)",
                "Mission-critical delivery, documentation, repeatable deployments, and compliance-aligned secure configuration support.",
                "",
                "Independent Consulting — Software Engineer / Security & Automation Consultant (Various)",
                "Client delivery, automation workflows, and documentation-first implementations.",
            ]),
            ("Featured Projects", [
                "Portable Translator — Offline OCR + Translation Services (github.com/brianhannigan/translator)",
                *translator_lines,
            ]),
            ("Education", education_lines_short()),
            ("Certifications", certification_lines()),
        ]
        return summary, skills, sections

    if variant == "python-backend":
        summary = "Backend-focused Python Engineer experienced building RESTful APIs, modular services, and data-driven workflows."
        skills = [
            "Python: APIs, services, tooling, scripting",
            "Backend: REST, JSON, HTTP, validation, error handling",
            "Databases: SQLite, SQL fundamentals, relational modeling",
            "Tools: Git, Docker, Windows/Linux",
            "Engineering: OOP, data structures, debugging, logging",
        ]
        sections = [
            ("Experience", [
                "Log(N) Pacific — Cyber Security Support Analyst (Vulnerability Management & SecOps Intern) (2025–Present)",
                "Tenable-based vulnerability scanning and remediation reporting; DISA STIG audits; PowerShell automation.",
                "EDR hunting, Sentinel dashboards, KQL investigations, and NSG/firewall hardening.",
                "",
                "Independent Consulting — Software Engineer / Security & Automation Consultant (Various)",
                "Python utilities and API-driven workflows; documentation-first delivery.",
                "",
                "UTRS, Inc — Technical Systems Engineer / Software Developer (2011–2025)",
                "Automation tooling, repeatable deployment procedures, disciplined troubleshooting, sustainment documentation.",
            ]),
            ("Featured Python Project", [
                "Portable Translator — Python Backend & API Services (github.com/brianhannigan/translator)",
                *translator_lines,
            ]),
            ("Education", education_lines_short()),
            ("Certifications", certification_lines()),
        ]
        return summary, skills, sections

    if variant == "cybersecurity":
        summary = (
            "Cybersecurity Engineer and systems builder with hands-on vulnerability management and SecOps internship experience, "
            "plus long-term secure systems delivery in mission-critical environments."
        )
        skills = [
            "Vulnerability Management: Tenable, risk prioritization, remediation reporting, Windows/Linux exposure",
            "Security Operations: EDR hunting, MDE detections, tabletop scenario design, incident triage concepts",
            "SIEM: Sentinel dashboards, KQL investigations, threat intelligence enrichment concepts",
            "Hardening: NSG/firewall rules, exposure reduction, DISA STIG audit concepts",
            "Documentation: workpapers, evidence organization, audit-ready reporting",
        ]
        sections = [
            ("Experience", [
                "Log(N) Pacific — Cyber Security Support Analyst (Vulnerability Management & SecOps Intern) (2025–Present)",
                "Vulnerability scans + remediation reporting; 100% reduction in critical, 90% in high, 76% in medium vulnerabilities.",
                "Tenable risk prioritization across Windows/Linux; DISA STIG audits; PowerShell remediation automation.",
                "EDR hunting; MDE detection rules; Sentinel dashboards; KQL investigations; exposure hardening via NSG/firewall.",
                "",
                "UTRS, Inc — Secure Systems / Technical Systems Engineer (2011–2025)",
                "Secure configuration support, compliance-aligned deployments, documentation discipline, operational troubleshooting.",
                "",
                "Independent Consulting — Security & Automation Consultant (Various)",
                "Evidence-friendly documentation, remediation tracking templates, delivery support.",
            ]),
            ("Security Portfolio", [
                "KQL Threat Hunting Beginner Guide — github.com/brianhannigan/kql-threathunting-beginner-guide",
                "Portable Translator — local/offline services with logging and repeatable startup workflows.",
            ]),
            ("Education", education_lines_short()),
            ("Certifications", certification_lines()),
        ]
        return summary, skills, sections

    if variant == "ai-architect":
        summary = (
            "AI Systems Architect and offline-first builder focused on practical AI integration, deployable workflows, "
            "reliability, and documentation."
        )
        skills = [
            "Offline-first architecture, deterministic workflows, deployment repeatability",
            "APIs & Services: REST, JSON, modular boundaries, logging/observability",
            "Automation: scripting, repeatable deployments, configuration patterns",
            "Tools: Python, Git, Docker, Windows/Linux",
            "Enablement: training, SOPs, stakeholder briefings",
        ]
        sections = [
            ("Experience", [
                "Log(N) Pacific — Cyber Security Support Analyst (Vulnerability Management & SecOps Intern) (2025–Present)",
                "Security operations + VM workflows reinforcing reliability and evidence discipline.",
                "",
                "Independent Consulting — Software Engineer / Automation Consultant (Various)",
                "Workflow automation patterns designed for maintainability and operational handoff.",
                "",
                "UTRS, Inc — Technical Systems Architect (2011–2025)",
                "Mission-critical delivery with emphasis on sustainment documentation and repeatable deployments.",
            ]),
            ("Flagship Project", [
                "Portable Translator — Offline OCR + Translation Orchestration (github.com/brianhannigan/translator)",
                *translator_lines,
            ]),
            ("Education", education_lines_short()),
            ("Certifications", certification_lines()),
        ]
        return summary, skills, sections

    if variant == "grc-analyst":
        summary = (
            "GRC Analyst and compliance-focused security professional with experience supporting secure deployments, "
            "evidence-driven documentation, and risk-aware engineering."
        )
        skills = [
            "GRC: risk assessments, evidence collection, control mapping, gap tracking, remediation planning",
            "Assessment Support: stakeholder interviews, documentation review, audit trails, follow-up tracking",
            "Framework Exposure: HIPAA concepts, NIST concepts, SOC 2 concepts",
            "Technical Context: vulnerability remediation workflows, secure configuration, Windows/Linux exposure",
            "Communication: concise writing, executive-ready summaries, status reporting",
        ]
        sections = [
            ("Experience", [
                "Log(N) Pacific — Cyber Security Support Analyst (Vulnerability Management & SecOps Intern) (2025–Present)",
                "Evidence-friendly remediation reporting and control-aligned activities (Tenable scans, DISA STIG audits, automation).",
                "Measurable risk reduction outcomes (critical/high/medium reductions; exposure hardening).",
                "",
                "Independent Consulting — Security & Automation Consultant (Various)",
                "Structured trackers + documentation templates supporting audit readiness and follow-ups.",
                "",
                "UTRS, Inc — Technical Systems Engineer (2011–2025)",
                "SOPs, sustainment artifacts, deployment documentation, compliance-aligned operational practices.",
            ]),
            ("Projects", [
                "Portable Translator — reliable local services with logging and repeatable startup workflows.",
            ]),
            ("Education", education_lines_short()),
            ("Certifications", certification_lines()),
        ]
        return summary, skills, sections

    raise ValueError(f"Unknown variant: {variant}")


# =========================
# MADAKOR FORMAT CONTENT + RENDERERS
# =========================
def madakor_blocks_for_variant(variant: Variant, cfg: ResumeConfig):
    variant = _fallback_variant(variant)

    experience_blocks = [
        logn_pacific_block(),
        consulting_block_for_variant(variant),
        utrs_block_for_variant(variant),
    ]

    proj = translator_project_for_variant(variant)
    certs = certification_lines()
    edu = education_lines_short()

    additional_common = [
        "Endpoint Detection and Response (concepts)", "Vulnerability Management", "CVE/CWE Concepts", "CVSS Concepts",
        "Risk Prioritization", "Remediation Tracking", "PowerShell Scripting", "Bash Scripting",
        "Firewall/NSG Configuration", "DISA STIG Alignment (audit concepts)", "NIST Concepts",
        "HIPAA Concepts", "SOC 2 Concepts", "Audit Trail Mindset",
        "REST APIs", "JSON", "HTTP", "Python", "SQL", "Git", "Docker",
        "Microsoft Sentinel (dashboards/queries)", "Microsoft Defender for Endpoint (detections)",
        "KQL (investigation/detection queries)",
    ]

    if variant == "python-backend":
        additional = additional_common + ["API Design", "Request Validation", "Logging/Observability", "Service Boundaries", "OOP", "Data Structures"]
    elif variant == "cybersecurity":
        additional = additional_common + ["Threat Hunting", "Detection Engineering", "Incident Triage", "Threat Intelligence Enrichment"]
    elif variant == "grc-analyst":
        additional = additional_common + ["Control Mapping", "Gap Logs", "Evidence Collection", "Workpapers", "Remediation Plans"]
    elif variant == "ai-architect":
        additional = additional_common + ["Offline-first Architecture", "Workflow Orchestration", "Reliability Engineering", "Configuration Management"]
    else:
        additional = additional_common + ["Systems Engineering", "Documentation", "Stakeholder Communication"]

    return experience_blocks, proj, certs, edu, additional


def render_classic_docx(doc: Document, cfg: ResumeConfig, summary: str, skills: List[str], sections: List[Section]) -> None:
    add_title(doc, cfg.name)
    add_center_line(doc, f"{cfg.location} | {cfg.phone} | {cfg.email}")
    add_center_line(doc, f"LinkedIn: {cfg.linkedin} | GitHub: {cfg.github}")

    add_section(doc, "Professional Summary")
    add_paragraph(doc, summary)

    add_section(doc, "Core Skills")
    add_bullets(doc, skills)

    for title, body in sections:
        add_section(doc, title)
        if isinstance(body, str):
            add_paragraph(doc, body)
        else:
            add_bullets(doc, body)


def render_madakor_docx(
    doc: Document,
    cfg: ResumeConfig,
    experience_blocks,
    projects_block,
    certs: List[str],
    education: List[str],
    additional_skills: List[str],
) -> None:
    add_title(doc, cfg.name)
    add_center_line(doc, cfg.phone)
    add_center_line(doc, cfg.email)
    add_center_line(doc, cfg.github)
    add_center_line(doc, cfg.linkedin)

    add_section(doc, "Experience")
    for role in experience_blocks:
        add_bold_line(doc, f"Company: {role['company']}\t{role['dates']}")
        add_bold_line(doc, f"Title: {role['title']}")
        for area_name, bullets in role.get("areas", []):
            add_bold_line(doc, f"{area_name}:")
            add_bullets(doc, bullets)

    add_section(doc, "Projects")
    add_bold_line(doc, projects_block["title"])
    add_paragraph(doc, f"Source: {projects_block['source']}")
    add_paragraph(doc, f"Platforms and Technology Used: {projects_block['platforms']}")
    add_bullets(doc, projects_block.get("bullets", []))

    add_section(doc, "Certifications")
    add_bullets(doc, certs)

    add_section(doc, "Education")
    add_bullets(doc, education)

    add_section(doc, "Additional Skills and Technologies")
    add_bullets(doc, additional_skills)


# =========================
# BUILD (DOCX/TXT/ODT/PDF)
# =========================
def build_variant(
    variant: Variant,
    cfg: ResumeConfig,
    also_odt: bool = False,
    also_pdf: bool = False,
) -> Dict[str, Path]:
    ensure_folders(cfg)

    # Route output to resumes/<variant>/<format>/
    fmt_folder = format_folder_name(cfg)
    out_dir = OUT / safe_slug(variant, default="general") / fmt_folder
    out_dir.mkdir(parents=True, exist_ok=True)

    slug = variant_slug(safe_slug(variant, default="general"))
    docx_path = out_dir / f"{slug}.docx"
    txt_path = out_dir / f"{slug}.txt"

    doc = Document()
    set_doc_margins(doc)

    if fmt_folder == "madakor":
        exp_blocks, proj_block, certs, edu, additional = madakor_blocks_for_variant(variant, cfg)
        render_madakor_docx(doc, cfg, exp_blocks, proj_block, certs, edu, additional)

        lines: List[str] = [
            cfg.name,
            cfg.phone,
            cfg.email,
            cfg.github,
            cfg.linkedin,
            "",
            "EXPERIENCE",
        ]
        for role in exp_blocks:
            lines.append(f"Company: {role['company']}\t{role['dates']}")
            lines.append(f"Title: {role['title']}")
            for area_name, bullets in role.get("areas", []):
                lines.append(f"{area_name}:")
                for b in bullets:
                    lines.append(f"- {b}")
            lines.append("")
        lines.append("PROJECTS")
        lines.append(proj_block["title"])
        lines.append(f"Source: {proj_block['source']}")
        lines.append(f"Platforms and Technology Used: {proj_block['platforms']}")
        for b in proj_block.get("bullets", []):
            lines.append(f"- {b}")
        lines.append("")
        lines.append("CERTIFICATIONS")
        for c in certs:
            lines.append(f"- {c}")
        lines.append("")
        lines.append("EDUCATION")
        for e in edu:
            lines.append(f"- {e}")
        lines.append("")
        lines.append("ADDITIONAL SKILLS AND TECHNOLOGIES")
        for s in additional:
            lines.append(f"- {s}")

        save_txt(txt_path, lines)

    else:
        summary, skills, sections = classic_content_variant(variant)
        render_classic_docx(doc, cfg, summary, skills, sections)

        lines: List[str] = []
        lines.extend([
            cfg.name,
            f"{cfg.location} | {cfg.phone} | {cfg.email}",
            f"LinkedIn: {cfg.linkedin} | GitHub: {cfg.github}",
            "",
            "PROFESSIONAL SUMMARY",
            summary,
            "",
            "CORE SKILLS",
        ])
        for s in skills:
            lines.append(f"- {s}")
        lines.append("")

        for title, body in sections:
            lines.append(title.upper())
            if isinstance(body, str):
                lines.append(body)
            else:
                for b in body:
                    lines.append(f"- {b}")
            lines.append("")

        save_txt(txt_path, lines)

    doc.save(str(docx_path))

    result: Dict[str, Path] = {"docx": docx_path, "txt": txt_path}

    if also_odt:
        result["odt"] = convert_with_libreoffice(docx_path, "odt", cfg)

    if also_pdf:
        result["pdf"] = convert_with_libreoffice(docx_path, "pdf", cfg)

    return result
import json
import argparse
from pathlib import Path
from datetime import datetime
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch

ROOT = Path(__file__).resolve().parents[1]
BLOCKS = ROOT / "blocks"
OUT_DEFAULT = Path(__file__).resolve().parent / "output"
BULLET_PREFIXES = ("-", "*", "•", "â€¢")

def load_profile(name):
    path = Path(__file__).resolve().parent / "profiles" / f"{name}.json"
    return json.loads(path.read_text(encoding="utf-8-sig"))

def read_block(path):
    p = ROOT / path
    return p.read_text(encoding="utf-8-sig", errors="ignore")

BULLET_PREFIXES = ("-", "*", "•", "â€¢")

IGNORE_SUBSTRINGS = (
    "[add bullets here]",
    "dates:**",
    "tools",
)

def clean_text(s: str) -> str:
    # Fix mojibake dashes
    s = s.replace("â€”", "—").replace("â€“", "–")
    return s.strip()

def extract_bullets(text: str):
    bullets = []
    for line in text.splitlines():
        s = clean_text(line)
        if not s:
            continue

        s_lower = s.lower()

        # Ignore placeholders / template noise
        if any(x in s_lower for x in IGNORE_SUBSTRINGS):
            continue

        # Only take bullet lines
        if s.startswith(BULLET_PREFIXES):
            item = s.lstrip("-*•â€¢ ").strip()
            item = clean_text(item)
            if item and not any(x in item.lower() for x in IGNORE_SUBSTRINGS):
                bullets.append(item)
        
    if not bullets:
        bullets = ["(No bullet content found — update blocks/experience/*.md)"]

    return bullets

def build_docx(profile, sections, out):
    doc = Document()
    doc.add_heading("Brian Hannigan", level=1)
    doc.add_paragraph(profile["summary"])

    doc.add_heading("Experience", level=2)
    for sec in sections:
        doc.add_heading(sec["title"], level=3)
        for b in sec["bullets"]:
            doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Skills", level=2)
    doc.add_paragraph(", ".join(profile["skills"]))
    doc.save(str(out))

def build_pdf(profile, sections, out):
    styles = getSampleStyleSheet()
    pdf = SimpleDocTemplate(str(out))
    elements = []

    elements.append(Paragraph("<b>Brian Hannigan</b>", styles["Heading1"]))
    elements.append(Paragraph(profile["summary"], styles["Normal"]))
    elements.append(Spacer(1, 0.2 * inch))

    elements.append(Paragraph("<b>Experience</b>", styles["Heading2"]))
    for sec in sections:
        elements.append(Paragraph(f"<b>{sec['title']}</b>", styles["Normal"]))
        for b in sec["bullets"]:
            elements.append(Paragraph(f"- {b}", styles["Normal"]))
        elements.append(Spacer(1, 0.15 * inch))

    elements.append(Paragraph("<b>Skills</b>", styles["Heading2"]))
    elements.append(Paragraph(", ".join(profile["skills"]), styles["Normal"]))

    pdf.build(elements)

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("profile")
    parser.add_argument("--outdir", default=str(OUT_DEFAULT))
    args = parser.parse_args()

    profile = load_profile(args.profile)
    sections = []

    for block in profile["experience_blocks"]:
        text = read_block(block["block"])
        bullets = extract_bullets(text)[:block.get("max_bullets", 6)]
        sections.append({
            "title": block["title"],
            "bullets": bullets
        })

    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)

    year = datetime.now().year
    base = f"Brian_Hannigan_{profile['output_name']}_{year}"
    docx_out = outdir / f"{base}.docx"
    pdf_out = outdir / f"{base}.pdf"

    build_docx(profile, sections, docx_out)
    build_pdf(profile, sections, pdf_out)

    print("Generated:")
    print(docx_out)
    print(pdf_out)

if __name__ == "__main__":
    main()
